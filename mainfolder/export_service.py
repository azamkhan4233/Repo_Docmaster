# ==============================
# export_service.py
# ==============================

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from PIL import Image
import os

# ==============================
# STYLE PRESETS
# ==============================
STYLE_PRESETS = {
    "APA": {
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 2.0,
        "margin_in": 1.0,
        "heading_size": 14,
        "alignment": "left",   # paragraph alignment
        "title_page": True
    },
    "IEEE": {
        "font": "Times New Roman",
        "font_size": 10,
        "line_spacing": 1.0,
        "margin_in": 0.75,
        "heading_size": 12,
        "alignment": "justify",
        "title_page": False
    },
    "MLA": {
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 2.0,
        "margin_in": 1.0,
        "heading_size": 13,
        "alignment": "left",
        "title_page": True
    },
    "Custom": {
        "font": "Arial",
        "font_size": 12,
        "line_spacing": 1.5,
        "margin_in": 1.0,
        "heading_size": 14,
        "alignment": "left",
        "title_page": True
    }
}


class ExportService:
    """
    Takes structured sections + style and exports to DOCX or PDF.
    Images: fixed width 4.5 inches, proportional height.
    """

    def __init__(self, style_choice="APA", custom_style=None, author=None, institution=None):
        self.style = STYLE_PRESETS.get(style_choice, STYLE_PRESETS["APA"]).copy()
        if custom_style:
            self.style.update(custom_style)

        self.author = author or "Anonymous"
        self.institution = institution or "Unknown Institution"

        # image width in inches (W3)
        self.image_width_in = 4.5

    # ==============================
    # DOCX EXPORT
    # ==============================
    def export_to_docx(self, sections, title, out_path):
        doc = Document()

        # Page margins
        margin = float(self.style.get("margin_in", 1.0))
        for s in doc.sections:
            s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(margin)

        # Normal style defaults
        try:
            normal = doc.styles["Normal"]
            normal.font.name = self.style["font"]
            normal.font.size = Pt(self.style["font_size"])
        except Exception:
            pass

        # Optional title page
        if self.style.get("title_page", True):
            self._add_title_page(doc, title)

        # Sections
        for sec_name, elems in sections.items():
            if not elems:
                continue

            # Heading (always left aligned)
            h = doc.add_heading(sec_name, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in h.runs:
                r.font.name = self.style["font"]
                r.font.size = Pt(self.style["heading_size"])
                r.bold = True

            # Content
            for el in elems:
                t = el.get("type")
                if t == "text":
                    self._add_paragraph(doc, el.get("content", ""))
                elif t == "table":
                    self._add_table(doc, el.get("data", []))
                elif t == "image":
                    self._add_image(doc, el)

        doc.save(out_path)

    def _add_title_page(self, doc, title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.name = self.style["font"]
        r.font.size = Pt(self.style["heading_size"] + 4)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Author: {self.author}\n")
        meta.add_run(f"Institution: {self.institution}\n")
        meta.add_run(datetime.now().strftime("%B %d, %Y"))

        doc.add_page_break()

    def _add_paragraph(self, doc, text):
        p = doc.add_paragraph(text)
        align = self.style.get("alignment", "left").lower()
        if align == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p.paragraph_format.line_spacing = self.style["line_spacing"]
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = self.style["font"]
            r.font.size = Pt(self.style["font_size"])

    def _add_table(self, doc, rows):
        if not rows:
            return
        table = doc.add_table(rows=0, cols=max(len(r) for r in rows))
        table.style = "Table Grid"  # T-grid
        for row in rows:
            cells = table.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)

    def _add_image(self, doc, image_dict):
        path = image_dict.get("path")
        caption = image_dict.get("caption")
        virtual = image_dict.get("virtual", False)

        if virtual:
            # Insert placeholder text
            p = doc.add_paragraph("[GRAPH/CHART PLACEHOLDER]")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = self.style["font"]
                r.font.size = Pt(self.style["font_size"])
            if caption:
                self._add_paragraph(doc, caption)
            return

        if path and os.path.exists(path):
            try:
                # Fixed width, proportional height (S2 + W3)
                doc.add_picture(path, width=Inches(self.image_width_in))
            except Exception:
                doc.add_paragraph(f"[Image could not be rendered: {path}]")
        else:
            doc.add_paragraph("[Missing Image Data]")

        if caption:
            self._add_paragraph(doc, caption)

    # ==============================
    # PDF EXPORT (Text + images + captions)
    # ==============================
    def export_to_pdf(self, sections, title, out_path):
        c = canvas.Canvas(out_path, pagesize=A4)
        w, h = A4
        y = h - 60

        # Title
        c.setFont("Times-Bold", 16)
        c.drawCentredString(w / 2, y, title)
        y -= 30

        c.setFont("Times-Roman", 11)
        c.drawCentredString(w / 2, y, f"{self.author} | {self.institution}")
        y -= 40

        max_img_width = self.image_width_in * 72  # inches -> points

        for sec, elems in sections.items():
            if y < 80:
                c.showPage()
                y = h - 60

            c.setFont("Times-Bold", 14)
            c.drawString(50, y, sec)
            y -= 20

            for el in elems:
                t = el.get("type")

                if t == "text":
                    c.setFont("Times-Roman", 11)
                    for line in self._wrap(el.get("content", ""), 90):
                        if y < 60:
                            c.showPage()
                            y = h - 60
                        c.drawString(50, y, line)
                        y -= 14 * self.style.get("line_spacing", 1.5)

                elif t == "image":
                    caption = el.get("caption")
                    virtual = el.get("virtual", False)
                    path = el.get("path")

                    if virtual:
                        c.setFont("Times-Italic", 10)
                        if y < 70:
                            c.showPage()
                            y = h - 60
                        c.drawString(50, y, "[GRAPH/CHART PLACEHOLDER]")
                        y -= 16
                        if caption:
                            c.drawString(50, y, caption[:110])
                            y -= 18
                        continue

                    if path and os.path.exists(path):
                        try:
                            # Use PIL to compute proportional height
                            with Image.open(path) as im:
                                iw, ih = im.size
                            scale = max_img_width / float(iw)
                            img_w = max_img_width
                            img_h = ih * scale

                            if y - img_h < 60:
                                c.showPage()
                                y = h - 60

                            c.drawImage(path, 50, y - img_h, width=img_w, height=img_h, preserveAspectRatio=True, mask='auto')
                            y -= img_h + 10

                            if caption:
                                c.setFont("Times-Italic", 10)
                                c.drawString(50, y, caption[:110])
                                y -= 18
                        except Exception:
                            c.setFont("Times-Italic", 10)
                            if y < 70:
                                c.showPage()
                                y = h - 60
                            c.drawString(50, y, f"[Image could not be rendered: {path}]")
                            y -= 16
                            if caption:
                                c.drawString(50, y, caption[:110])
                                y -= 18

        c.save()

    def _wrap(self, text, width=90):
        words = text.split()
        out = []
        line = ""
        for w in words:
            if len(line) + len(w) + 1 <= width:
                line = (line + " " + w).strip()
            else:
                out.append(line)
                line = w
        if line:
            out.append(line)
        return out

# ==============================
# export_service.py (Improved)
# ==============================

# from docx import Document
# from docx.shared import Inches, Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from datetime import datetime
# from reportlab.lib.pagesizes import A4
# from reportlab.pdfgen import canvas
# from PIL import Image
# import os

# # ==============================
# # STYLE PRESETS
# # ==============================
# STYLE_PRESETS = {
#     "APA": {
#         "font": "Times New Roman", "font_size": 12, "line_spacing": 2.0,
#         "margin_in": 1.0, "heading_size": 14, "alignment": "left", "title_page": True
#     },
#     "IEEE": {
#         "font": "Times New Roman", "font_size": 10, "line_spacing": 1.0,
#         "margin_in": 0.75, "heading_size": 12, "alignment": "justify", "title_page": False
#     },
#     "MLA": {
#         "font": "Times New Roman", "font_size": 12, "line_spacing": 2.0,
#         "margin_in": 1.0, "heading_size": 13, "alignment": "left", "title_page": True
#     },
#     "Custom": {
#         "font": "Arial", "font_size": 12, "line_spacing": 1.5,
#         "margin_in": 1.0, "heading_size": 14, "alignment": "left", "title_page": True
#     }
# }


# class ExportService:
#     """Export formatted DOCX + PDF preserving order + captions."""

#     def __init__(self, style_choice="APA", custom_style=None, author=None, institution=None):
#         self.style = STYLE_PRESETS.get(style_choice, STYLE_PRESETS["APA"]).copy()
#         if custom_style:
#             self.style.update(custom_style)

#         self.author = author or "Anonymous"
#         self.institution = institution or "Unknown Institution"
#         self.image_width_in = 4.5  # (W3 user requirement)

#     # ==============================
#     # DOCX EXPORT
#     # ==============================
#     def export_to_docx(self, sections, title, out_path):
#         doc = Document()

#         # Page margins
#         margin = float(self.style["margin_in"])
#         for s in doc.sections:
#             s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(margin)

#         # Default normal style
#         try:
#             normal = doc.styles["Normal"]
#             normal.font.name = self.style["font"]
#             normal.font.size = Pt(self.style["font_size"])
#         except Exception:
#             pass

#         # Title page
#         if self.style.get("title_page", True):
#             self._add_title_page(doc, title)

#         # Body
#         for sec_name, elems in sections.items():
#             if not elems:
#                 continue

#             self._add_section_heading(doc, sec_name)

#             for el in elems:
#                 match el.get("type"):
#                     case "text": self._add_paragraph(doc, el["content"])
#                     case "table": self._add_table(doc, el["data"])
#                     case "image": self._add_image(doc, el)

#         doc.save(out_path)

#     # -------- SECTION HEADING --------
#     def _add_section_heading(self, doc, sec):
#         h = doc.add_heading(sec, level=1)
#         h.alignment = WD_ALIGN_PARAGRAPH.LEFT
#         for r in h.runs:
#             r.font.name = self.style["font"]
#             r.font.size = Pt(self.style["heading_size"])
#             r.bold = True

#     # -------- TITLE PAGE --------
#     def _add_title_page(self, doc, title):
#         p = doc.add_paragraph()
#         p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         r = p.add_run(title)
#         r.bold = True
#         r.font.name = self.style["font"]
#         r.font.size = Pt(self.style["heading_size"] + 4)

#         meta = doc.add_paragraph()
#         meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         meta.add_run(f"Author: {self.author}\n")
#         meta.add_run(f"Institution: {self.institution}\n")
#         meta.add_run(datetime.now().strftime("%B %d, %Y"))
#         doc.add_page_break()

#     # -------- TEXT --------
#     def _add_paragraph(self, doc, text):
#         p = doc.add_paragraph(text)
#         p.paragraph_format.line_spacing = self.style["line_spacing"]
#         p.paragraph_format.space_after = Pt(6)

#         align = self.style.get("alignment", "left")
#         p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == "justify" else WD_ALIGN_PARAGRAPH.LEFT

#         for r in p.runs:
#             r.font.name = self.style["font"]
#             r.font.size = Pt(self.style["font_size"])

#     # -------- TABLE --------
#     def _add_table(self, doc, rows):
#         if not rows:
#             return
#         table = doc.add_table(rows=0, cols=max(len(r) for r in rows))
#         table.style = "Table Grid"
#         for row in rows:
#             cells = table.add_row().cells
#             for i, v in enumerate(row):
#                 cells[i].text = str(v)
#         doc.add_paragraph("")  # spacing

#     # -------- IMAGES + CAPTIONS --------
#     def _add_image(self, doc, image):
#         path, caption, virtual = image.get("path"), image.get("caption"), image.get("virtual", False)

#         # Virtual / missing images
#         if virtual or not (path and os.path.exists(path)):
#             p = doc.add_paragraph("[GRAPH/IMAGE PLACEHOLDER]")
#             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             return self._caption(doc, caption)

#         # Insert real image
#         try:
#             doc.add_picture(path, width=Inches(self.image_width_in))
#         except Exception:
#             doc.add_paragraph(f"[Image Error: {path}]")
#         return self._caption(doc, caption)

#     # -------- CAPTION STYLE --------
#     def _caption(self, doc, text):
#         if not text:
#             return
#         p = doc.add_paragraph(text)
#         p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         for r in p.runs:
#             r.font.italic = True
#             r.font.name = self.style["font"]
#             r.font.size = Pt(self.style["font_size"])

#     # ==============================
#     # PDF EXPORT
#     # ==============================
#     def export_to_pdf(self, sections, title, out_path):
#         c = canvas.Canvas(out_path, pagesize=A4)
#         w, h = A4
#         y = h - 60

#         # Title
#         c.setFont("Times-Bold", 16)
#         c.drawCentredString(w / 2, y, title); y -= 30
#         c.setFont("Times-Roman", 11)
#         c.drawCentredString(w / 2, y, f"{self.author} | {self.institution}"); y -= 40

#         max_img_width = self.image_width_in * 72

#         for sec, elems in sections.items():
#             if y < 80: c.showPage(); y = h - 60
#             c.setFont("Times-Bold", 14); c.drawString(50, y, sec); y -= 20

#             for el in elems:
#                 t = el.get("type")
#                 if t == "text": y = self._pdf_text(c, el["content"], y)
#                 elif t == "image": y = self._pdf_image(c, el, y, max_img_width)

#         c.save()

#     # ---- PDF HELPERS ----
#     def _pdf_text(self, c, text, y):
#         ln = self.style.get("line_spacing", 1.5)
#         c.setFont("Times-Roman", 11)
#         for line in self._wrap(text, 90):
#             if y < 60:
#                 c.showPage(); y = A4[1] - 60
#             c.drawString(50, y, line); y -= 14 * ln
#         return y - 6

#     def _pdf_image(self, c, img, y, max_w):
#         caption, virtual, path = img.get("caption"), img.get("virtual", False), img.get("path")

#         if virtual or not (path and os.path.exists(path)):
#             return self._pdf_caption(c, "[GRAPH/IMAGE PLACEHOLDER]", y)

#         try:
#             with Image.open(path) as im:
#                 iw, ih = im.size
#             scale = max_w / iw; ih *= scale
#             if y - ih < 60: c.showPage(); y = A4[1] - 60
#             c.drawImage(path, 50, y - ih, width=max_w, height=ih); y -= ih + 10
#         except:
#             return self._pdf_caption(c, "[Image Error]", y)

#         return self._pdf_caption(c, caption, y)

#     def _pdf_caption(self, c, text, y):
#         if not text: return y
#         c.setFont("Times-Italic", 10)
#         if y < 60: c.showPage(); y = A4[1] - 60
#         c.drawString(50, y, text[:110]); return y - 18

#     def _wrap(self, text, width=90):
#         words, out, line = text.split(), [], ""
#         for w in words:
#             if len(line) + len(w) + 1 <= width: line = (line + " " + w).strip()
#             else: out.append(line); line = w
#         if line: out.append(line); return out
